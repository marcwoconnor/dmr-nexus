#!/usr/bin/env python3
"""Generate a formatted Word document from the clustering TODO checklist."""

import re
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)


def add_formatted_text(paragraph, text):
    """Add text with inline bold/italic/code markdown formatting."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            paragraph.add_run(part)


def add_checkbox(paragraph, checked=False):
    """Add a checkbox character to the start of a paragraph."""
    symbol = '\u2611' if checked else '\u2610'  # ballot box with/without check
    run = paragraph.add_run(symbol + '  ')
    run.font.size = Pt(12)
    return run


def parse_markdown_to_docx(doc, md_text):
    """Parse markdown text and add it to a Word document with formatting."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['No Spacing']
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                pf = p.paragraph_format
                pf.space_before = Pt(4)
                pf.space_after = Pt(4)
                pf.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            if re.match(r'^[\s|:-]+$', line):
                i += 1
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append(cells)
            i += 1
            if i >= len(lines) or not (lines[i].strip().startswith('|')):
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                p.style = doc.styles['No Spacing']
                                add_formatted_text(p, cell_text)
                                p.paragraph_format.space_before = Pt(2)
                                p.paragraph_format.space_after = Pt(2)
                                for run in p.runs:
                                    run.font.size = Pt(8.5)

                                if row_idx == 0:
                                    set_cell_shading(cell, '2E4057')
                                    for run in p.runs:
                                        run.bold = True
                                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                        run.font.size = Pt(8.5)
                                elif row_idx % 2 == 0:
                                    set_cell_shading(cell, 'F2F2F2')

                    doc.add_paragraph()
                    table_rows = []
                    in_table = False
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            i += 1  # skip top-level (handled as title)
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Checkbox items: - [ ] or - [x]
        checkbox_match = re.match(r'^(\s*)- \[([ x])\]\s+(.*)', line)
        if checkbox_match:
            indent_str, check, text = checkbox_match.groups()
            indent_level = len(indent_str) // 2  # 2 spaces per indent level
            checked = check == 'x'

            p = doc.add_paragraph()
            p.style = doc.styles['No Spacing']
            # Base indent + extra for nesting
            p.paragraph_format.left_indent = Inches(0.4 + indent_level * 0.3)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            add_checkbox(p, checked)
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list (with possible sub-items)
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if numbered_match:
            indent_str, text = numbered_match.groups()
            indent_level = len(indent_str) // 2
            p = doc.add_paragraph(style='List Number')
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.5 + indent_level * 0.3)
            add_formatted_text(p, text)
            i += 1
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line.strip())
        i += 1


def create_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Customize heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    h1 = doc.styles['Heading 1']
    h1.font.size = Pt(22)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles['Heading 2']
    h2.font.size = Pt(16)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.page_break_before = True  # each phase starts new page

    h3 = doc.styles['Heading 3']
    h3.font.size = Pt(13)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    h4 = doc.styles['Heading 4']
    h4.font.size = Pt(11)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)

    # ===== TITLE PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('HBlink4', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_heading('Clustering Implementation TODO', level=0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.bold = False

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        'Detailed checklist covering Phases 1-6:\n'
        'Cluster Bus, Cross-Server Routing, Failover,\n'
        'Operational Tools, Native Protocol, and Global Scale'
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'Generated: {date.today().isoformat()}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_page_break()

    # ===== CONTENT =====
    with open('clustering_todo.md', 'r') as f:
        content = f.read()

    # Remove the H1 title line (we have our own title page)
    content = re.sub(r'^# .*\n', '', content)
    # Remove the "Generated" line at the bottom (we have it on the title page)
    content = re.sub(r'\n\*Generated:.*\*\s*$', '', content)

    parse_markdown_to_docx(doc, content)

    # Save
    output_path = 'HBlink4_Clustering_TODO.docx'
    doc.save(output_path)
    print(f'Generated: {output_path}')


if __name__ == '__main__':
    create_document()
