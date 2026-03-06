#!/usr/bin/env python3
"""Generate a formatted Word document from the HBlink4 design documentation."""

import re
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)


def add_formatted_text(paragraph, text):
    """Add text with inline bold/italic/code markdown formatting."""
    # Process inline formatting: **bold**, *italic*, `code`
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            paragraph.add_run(part)


def parse_markdown_to_docx(doc, md_text, is_first_section=False):
    """Parse markdown text and add it to a Word document with formatting."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - add collected lines
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['No Spacing']
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                pf = p.paragraph_format
                pf.space_before = Pt(4)
                pf.space_after = Pt(4)
                pf.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            # Check if this is a separator line (|---|---|)
            if re.match(r'^[\s|:-]+$', line):
                i += 1
                continue
            # Parse table row
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append(cells)
            i += 1
            # Check if next line is end of table
            if i >= len(lines) or not (lines[i].strip().startswith('|')):
                # Render table
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                p.style = doc.styles['No Spacing']
                                add_formatted_text(p, cell_text)
                                p.paragraph_format.space_before = Pt(2)
                                p.paragraph_format.space_after = Pt(2)
                                for run in p.runs:
                                    run.font.size = Pt(8.5)

                                # Header row styling
                                if row_idx == 0:
                                    set_cell_shading(cell, '2E4057')
                                    for run in p.runs:
                                        run.bold = True
                                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                        run.font.size = Pt(8.5)
                                elif row_idx % 2 == 0:
                                    set_cell_shading(cell, 'F2F2F2')

                    doc.add_paragraph()  # spacing after table
                    table_rows = []
                    in_table = False
            continue

        # Horizontal rule
        if line.strip() == '---':
            # Add a subtle separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Blockquote
        if line.strip().startswith('> '):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Headings
        if line.startswith('# ') and not is_first_section:
            # Top-level heading - treated as a section title within the doc
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('# '):
            # Skip - will be handled as page title
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line.strip())
        i += 1


def create_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Customize heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    h1 = doc.styles['Heading 1']
    h1.font.size = Pt(22)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles['Heading 2']
    h2.font.size = Pt(16)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles['Heading 3']
    h3.font.size = Pt(13)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    h4 = doc.styles['Heading 4']
    h4.font.size = Pt(11)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)

    # ===== TITLE PAGE =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('HBlink4', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_heading('Architecture, Clustering, and Protocol Design', level=0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.bold = False

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('DMR Server Implementation using the HomeBrew Protocol')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()
    doc.add_paragraph()

    toc_intro = doc.add_paragraph()
    toc_intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_intro.add_run('This document covers four areas:')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    items = [
        'Part 1: Operational Overview — How HBlink4 works today',
        'Part 2: Clustering Plan — Multi-server design with 5 implementation phases',
        'Part 3: Cluster-Native Protocol — New client protocol for fast failover',
        'Part 4: Global-Scale Architecture — Hierarchical routing for worldwide deployment',
    ]
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ===== REVISION HISTORY =====
    doc.add_heading('Revision History', level=2)
    doc.add_paragraph()

    revisions = [
        ['Version', 'Date', 'Author', 'Description'],
        ['1.0', '2026-03-06', 'Claude Opus 4.6', 'Initial document: Operational Overview, Clustering Plan (Phases 1-5), Cluster-Native Protocol Design'],
        ['1.1', '2026-03-06', 'Claude Opus 4.6', 'Added risk mitigations (asyncio load, stream ID collisions, config consistency) based on design review'],
        ['1.2', '2026-03-06', 'Claude Opus 4.6', 'Added Part 4: Global-Scale Architecture with hierarchical routing, regional clusters, and backbone design'],
        ['1.3', '2026-03-06', 'Claude Opus 4.6', 'Addressed global-scale review: backbone loop prevention, gateway backpressure, wildcard TG semantics, negative caching, backbone security hardening (Phase 6.6)'],
    ]

    rev_table = doc.add_table(rows=len(revisions), cols=4)
    rev_table.style = 'Table Grid'
    rev_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    col_widths = [Inches(0.7), Inches(1.1), Inches(1.5), Inches(3.7)]
    for row_idx, row_data in enumerate(revisions):
        for col_idx, cell_text in enumerate(row_data):
            cell = rev_table.rows[row_idx].cells[col_idx]
            cell.width = col_widths[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.style = doc.styles['No Spacing']
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            if row_idx == 0:
                set_cell_shading(cell, '2E4057')
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph()
    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    doc.add_heading('Table of Contents', level=2)
    doc.add_paragraph()

    # Insert a Word TOC field (updates when user presses F9 or right-click > Update Field in Word)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar2)

    # Placeholder text shown before TOC is updated
    run4 = p.add_run('Right-click and select "Update Field" to generate table of contents.')
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run4.font.italic = True
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar3)

    doc.add_paragraph()

    # Also add a static quick reference TOC for readers who don't update fields
    doc.add_heading('Document Structure', level=3)

    toc_entries = [
        ('Part 1: Operational Overview', [
            'The Big Picture',
            'The Two Processes',
            'The Core Server Module by Module',
            'Configuration',
            'How to Actually Use It',
            'Key Concepts to Remember',
        ]),
        ('Part 2: Clustering Plan', [
            'Current State Analysis',
            'Design Principles',
            'Architecture (Cluster Topology, Server Identity)',
            'Phase 1: Cluster Bus and State Awareness',
            'Phase 2: Cross-Server Stream Routing',
            'Phase 3: Failover and Resilience',
            'Phase 4: Shared Configuration and Operational Tools',
            'Known Risks and Mitigations',
            'Implementation Order and Effort Estimates',
            'Key Design Decisions and Tradeoffs',
            'Phase 5: Cluster-Native Client Protocol',
            'Scaling Beyond 15 Nodes',
        ]),
        ('Part 3: Cluster-Native Protocol Design', [
            'What the HomeBrew Protocol Forces on Us',
            'What a Cluster-Native Protocol Would Look Like',
            'What Changes in HBlink4',
            'Migration Strategy',
            'Comparison: HomeBrew vs. Native Clustering',
            'Recommendation',
        ]),
        ('Part 4: Global-Scale Clustering Architecture', [
            'Why the Small-Cluster Design Breaks Down',
            'Hierarchical Routing with Regional Clusters',
            'Backbone Forwarding Rule',
            'Three Tiers of Communication',
            'Talkgroup Routing Table (Wildcard Semantics, Size Bounds)',
            'Gateway Overload and Degradation Policy',
            'User Cache at Global Scale (Negative Caching, Stale Handling)',
            'Scaling Math',
            'Implementation (Phase 6)',
            'Backbone Security Considerations',
        ]),
    ]

    for part_title, sections in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(part_title)
        run.bold = True
        run.font.size = Pt(11)
        for section in sections:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(section)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # ===== PART 1: OPERATIONAL OVERVIEW =====
    part1_title = doc.add_heading('Part 1', level=1)
    for run in part1_title.runs:
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(14)

    with open('operational_overview.md', 'r') as f:
        content = f.read()
    # Remove the H1 title (we added our own) and the Further Reading section
    content = re.sub(r'^# .*\n', '', content)
    content = re.sub(r'---\s*\n## Further Reading.*', '', content, flags=re.DOTALL)
    parse_markdown_to_docx(doc, content, is_first_section=True)

    doc.add_page_break()

    # ===== PART 2: CLUSTERING PLAN =====
    part2_title = doc.add_heading('Part 2', level=1)
    for run in part2_title.runs:
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(14)

    with open('clustering_plan.md', 'r') as f:
        content = f.read()
    content = re.sub(r'^# .*\n', '', content)
    # Remove markdown link syntax from Phase 5 reference
    content = content.replace(
        'See [Cluster-Native Protocol Design](cluster_native_protocol.md)',
        'See Part 3: Cluster-Native Protocol Design'
    )
    content = content.replace(
        'See [Global-Scale Clustering Architecture](global_scaling.md)',
        'See Part 4: Global-Scale Clustering Architecture'
    )
    parse_markdown_to_docx(doc, content, is_first_section=True)

    doc.add_page_break()

    # ===== PART 3: CLUSTER-NATIVE PROTOCOL =====
    part3_title = doc.add_heading('Part 3', level=1)
    for run in part3_title.runs:
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(14)

    with open('cluster_native_protocol.md', 'r') as f:
        content = f.read()
    content = re.sub(r'^# .*\n', '', content)
    # Clean up markdown link in prerequisite note
    content = content.replace(
        '[Clustering Plan](clustering_plan.md)',
        'Part 2: Clustering Plan'
    )
    parse_markdown_to_docx(doc, content, is_first_section=True)

    doc.add_page_break()

    # ===== PART 4: GLOBAL-SCALE ARCHITECTURE =====
    part4_title = doc.add_heading('Part 4', level=1)
    for run in part4_title.runs:
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(14)

    with open('global_scaling.md', 'r') as f:
        content = f.read()
    content = re.sub(r'^# .*\n', '', content)
    # Clean up markdown links
    content = content.replace(
        '[Clustering Plan](clustering_plan.md)',
        'Part 2: Clustering Plan'
    )
    content = content.replace(
        'See [Global-Scale Clustering Architecture](global_scaling.md)',
        'See Part 4: Global-Scale Clustering Architecture'
    )
    parse_markdown_to_docx(doc, content, is_first_section=True)

    # Save
    output_path = 'HBlink4_Architecture_and_Clustering.docx'
    doc.save(output_path)
    print(f'Generated: {output_path}')


if __name__ == '__main__':
    create_document()
